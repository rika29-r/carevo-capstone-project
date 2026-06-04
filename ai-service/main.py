import os
os.environ.setdefault("TF_USE_LEGACY_KERAS", "1")

import json
import pickle
import re
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional

import numpy as np
import tensorflow as tf
from fastapi import FastAPI
from fastapi.responses import FileResponse
from pydantic import BaseModel
from career_rules import hybrid_recommendation

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except Exception:
    Document = None

BASE_DIR = Path(__file__).resolve().parent
MODEL_DIR = BASE_DIR / "model"
MODEL_PATH = MODEL_DIR / "career_model.keras"
LABEL_PATHS = [MODEL_DIR / "label_encoder.pkl", MODEL_DIR / "label_encode.pkl", BASE_DIR / "label_encoder.pkl", BASE_DIR / "label_encode.pkl"]
FALLBACK_LABELS_PATH = BASE_DIR / "labels.json"

MAX_TOKENS = 10000
MAX_LEN = 100
EMBEDDING_DIM = 64

app = FastAPI(title="CAREVO AI Service", version="1.7.0")

model: Optional[tf.keras.Model] = None
label_encoder = None
labels: List[str] = []
load_error = ""

PROFILE_TITLES = {
    "Administrasi": "Administrative Specialist",
    "Bisnis": "Business & HR Specialist",
    "Data & AI": "Data Science",
    "Keamanan Siber": "Cyber Security Analyst",
    "Kreatif & Desain": "UI/UX Designer",
    "Pemasaran": "Digital Marketing Specialist",
    "Pendidikan": "Education Specialist",
    "Rekayasa Perangkat Lunak": "Software Engineer",
}

TOP_PATHS = {
    "Administrasi": ["Administrative Officer", "Office Coordinator", "Data Entry Specialist"],
    "Bisnis": ["Human Resource Staff", "Business Analyst", "Project Coordinator"],
    "Data & AI": ["Data Scientist", "Data Analyst", "Machine Learning Engineer"],
    "Keamanan Siber": ["Cyber Security Analyst", "SOC Analyst", "Security Engineer"],
    "Kreatif & Desain": ["UI/UX Designer", "Product Designer", "Visual Designer"],
    "Pemasaran": ["Digital Marketing Specialist", "SEO Specialist", "Content Strategist"],
    "Pendidikan": ["Teacher", "Mentor", "Instructional Designer"],
    "Rekayasa Perangkat Lunak": ["Software Engineer", "Frontend Developer", "Backend Developer"],
}

KEYWORDS = {
    "Administrasi": {
        "administrasi": 7, "administration": 7, "administrative": 6, "admin": 4, "data entry": 8,
        "arsip": 7, "archive": 6, "filing": 5, "scheduling": 6, "penjadwalan": 6,
        "document management": 8, "manajemen dokumen": 8, "microsoft office": 5, "microsoft excel": 6,
        "excel": 5, "microsoft word": 4, "laporan": 4, "customer service": 4, "pelayanan pelanggan": 4,
    },
    "Bisnis": {
        "human resource": 12, "human resources": 12, "sumber daya manusia": 12, "hr": 8, "hrd": 9,
        "recruitment": 11, "rekrutmen": 11, "recruiter": 9, "candidate screening": 10,
        "screening kandidat": 10, "interview candidate": 9, "interview": 6, "wawancara": 6,
        "employee administration": 10, "administrasi karyawan": 10, "employee": 7, "karyawan": 7,
        "people management": 9, "talent acquisition": 9, "payroll": 7, "onboarding": 7,
        "business analysis": 8, "business analyst": 8, "business": 5, "bisnis": 5, "management": 6,
        "manajemen": 6, "project management": 8, "scrum": 5, "agile": 5, "strategy": 5,
        "financial planning": 7, "strategic management": 7, "leadership": 4, "communication": 3,
        "komunikasi": 3, "teamwork": 2, "public speaking": 2,
    },
    "Data & AI": {
        "data analysis": 11, "analisis data": 11, "data analyst": 11, "data science": 12,
        "data scientist": 12, "machine learning": 12, "deep learning": 10, "artificial intelligence": 10,
        "ai": 5, "tensorflow": 10, "pytorch": 9, "python": 6, "sql": 6, "pandas": 8,
        "numpy": 8, "statistics": 7, "statistik": 7, "data visualization": 7, "visualisasi data": 7,
        "big data": 7, "analytics": 6, "dashboard": 4, "tableau": 6, "powerbi": 6, "power bi": 6,
    },
    "Keamanan Siber": {
        "cybersecurity": 12, "cyber security": 12, "cyber": 8, "siber": 8, "keamanan siber": 12,
        "security": 5, "network security": 11, "penetration testing": 11, "pentest": 10,
        "ethical hacking": 10, "firewall": 8, "vulnerability": 8, "malware": 7, "linux": 5,
        "incident response": 8, "soc": 6, "security analyst": 9,
    },
    "Kreatif & Desain": {
        "ui ux": 12, "ux ui": 12, "ui/ux": 12, "ui": 5, "ux": 5, "figma": 11,
        "wireframe": 8, "wireframing": 8, "prototype": 8, "graphic design": 9, "desain grafis": 9,
        "design": 5, "desain": 5, "visual communication": 8, "komunikasi visual": 8,
        "adobe illustrator": 8, "illustrator": 7, "photoshop": 7, "creative": 5, "kreatif": 5, "branding": 4,
    },
    "Pemasaran": {
        "digital marketing": 12, "marketing": 10, "pemasaran": 10, "seo": 9, "sem": 8,
        "social media marketing": 9, "content strategy": 8, "campaign": 7, "google analytics": 7,
        "copywriting": 6, "advertising": 6, "ads": 5, "brand": 5, "branding": 5, "social": 3,
        "media": 3, "content": 5,
    },
    "Pendidikan": {
        "teaching": 11, "mengajar": 11, "teacher": 10, "guru": 10, "education": 9,
        "pendidikan": 9, "mentoring": 8, "mentor": 8, "curriculum": 9, "kurikulum": 9,
        "classroom management": 8, "educational assessment": 8, "assessment": 4, "training": 6,
        "tutor": 7, "pembelajaran": 6, "kelas": 5,
    },
    "Rekayasa Perangkat Lunak": {
        "software engineering": 12, "software engineer": 12, "rekayasa perangkat lunak": 12,
        "frontend": 11, "front end": 11, "backend": 11, "back end": 11, "fullstack": 11,
        "full stack": 11, "web developer": 9, "developer": 5, "programming": 9, "coding": 9,
        "javascript": 9, "typescript": 9, "react": 9, "vite": 5, "node js": 9, "node.js": 9,
        "express": 8, "rest api": 8, "api": 5, "microservices": 7, "spring boot": 8,
        "java": 5, "database": 3, "postgresql": 5, "mysql": 5, "docker": 6, "deployment": 5,
        "html": 3, "css": 3, "teknik informatika": 1, "computer science": 1, "informatika": 1,
    },
}

STRONG_HR = ["human resource", "human resources", "hrd", "sumber daya manusia", "recruitment", "rekrutmen", "candidate screening", "interview candidate", "employee administration", "administrasi karyawan", "people management", "talent acquisition"]
STRONG_SOFTWARE = ["frontend", "backend", "fullstack", "full stack", "react", "node js", "node.js", "express", "javascript", "typescript", "coding", "programming", "software engineering", "software engineer", "spring boot", "rest api", "microservices"]
STRONG_DATA = ["machine learning", "data analysis", "analisis data", "data science", "data scientist", "data analyst", "tensorflow", "pandas", "numpy", "statistics", "statistik"]

ALIASES = {
    "business": "Bisnis", "bisnis": "Bisnis", "bisnis & manajemen": "Bisnis", "human resource": "Bisnis", "hr": "Bisnis", "hrd": "Bisnis",
    "administration": "Administrasi", "administrasi": "Administrasi", "admin": "Administrasi",
    "data": "Data & AI", "data science": "Data & AI", "data & ai": "Data & AI", "ai": "Data & AI",
    "cybersecurity": "Keamanan Siber", "cyber security": "Keamanan Siber", "keamanan siber": "Keamanan Siber",
    "design": "Kreatif & Desain", "desain": "Kreatif & Desain", "ui ux": "Kreatif & Desain", "ui/ux": "Kreatif & Desain",
    "marketing": "Pemasaran", "pemasaran": "Pemasaran", "digital marketing": "Pemasaran",
    "education": "Pendidikan", "pendidikan": "Pendidikan", "teaching": "Pendidikan",
    "software": "Rekayasa Perangkat Lunak", "software engineering": "Rekayasa Perangkat Lunak", "software engineer": "Rekayasa Perangkat Lunak", "rekayasa perangkat lunak": "Rekayasa Perangkat Lunak",
}

class CareerRequest(BaseModel):
    skills_text: str = ""

class PredictCareerRequest(BaseModel):
    text: str = ""
    cvData: Dict[str, Any] = {}

class GenerateCvRequest(BaseModel):
    name: str = "CAREVO USER"
    phone: str = ""
    email: str = ""
    linkedin: str = ""
    portfolio: str = ""
    location: str = ""
    professional_summary: str = ""
    education: List[Dict[str, Any]] = []
    experiences: List[Dict[str, Any]] = []
    projects: List[Dict[str, Any]] = []
    certifications: List[str] = []
    achievements: List[str] = []
    skills: str = ""
    languages: str = ""
    career_group: str = ""
    recommended_careers: List[str] = []


def clean_text(text: Any) -> str:
    text = str(text or "").lower()
    text = re.sub(r"[^a-zA-Z0-9&+./# ]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def keyword_hit(text: str, keyword: str) -> bool:
    t = clean_text(text)
    k = clean_text(keyword)
    if not k:
        return False
    if len(k) <= 3 or k in ["ai", "hr", "ui", "ux", "qa"]:
        return re.search(rf"(^|\s){re.escape(k)}(\s|$)", t) is not None
    return k in t

def has_any(text: str, keywords: List[str]) -> bool:
    return any(keyword_hit(text, kw) for kw in keywords)

def normalize_category(value: str) -> str:
    text = clean_text(value)
    if not text:
        return ""
    if text in ALIASES:
        return ALIASES[text]
    for label in PROFILE_TITLES:
        if clean_text(label) == text:
            return label
    if keyword_hit(text, "human resource") or keyword_hit(text, "recruitment") or keyword_hit(text, "hrd"):
        return "Bisnis"
    if keyword_hit(text, "software") or keyword_hit(text, "frontend") or keyword_hit(text, "backend"):
        return "Rekayasa Perangkat Lunak"
    if keyword_hit(text, "machine learning") or keyword_hit(text, "data") or keyword_hit(text, "ai"):
        return "Data & AI"
    if keyword_hit(text, "cyber") or keyword_hit(text, "security") or keyword_hit(text, "siber"):
        return "Keamanan Siber"
    if keyword_hit(text, "figma") or keyword_hit(text, "design") or keyword_hit(text, "ui") or keyword_hit(text, "ux"):
        return "Kreatif & Desain"
    if keyword_hit(text, "marketing") or keyword_hit(text, "seo") or keyword_hit(text, "pemasaran"):
        return "Pemasaran"
    if keyword_hit(text, "teaching") or keyword_hit(text, "education") or keyword_hit(text, "pendidikan"):
        return "Pendidikan"
    if keyword_hit(text, "administrasi") or keyword_hit(text, "admin"):
        return "Administrasi"
    return value

def rule_scores(text: str) -> Dict[str, Any]:
    scores: Dict[str, int] = {}
    matched: Dict[str, List[str]] = {}
    for category, keywords in KEYWORDS.items():
        scores[category] = 0
        matched[category] = []
        for kw, weight in keywords.items():
            if keyword_hit(text, kw):
                scores[category] += weight
                matched[category].append(kw)
    if (keyword_hit(text, "teknik informatika") or keyword_hit(text, "informatika") or keyword_hit(text, "computer science")) and not has_any(text, STRONG_SOFTWARE):
        scores["Rekayasa Perangkat Lunak"] = min(scores.get("Rekayasa Perangkat Lunak", 0), 2)
    return {"scores": scores, "matched": matched}

def match_from_score(score: int) -> int:
    if score >= 34: return 97
    if score >= 26: return 94
    if score >= 18: return 90
    if score >= 12: return 84
    if score >= 8: return 78
    if score >= 4: return 68
    return 58

def best_rule(text: str):
    data = rule_scores(text)
    best = sorted(data["scores"].items(), key=lambda x: x[1], reverse=True)[0]
    return best, data

def recommendation(category: str, match_score: int, matched_keywords: List[str], extra: Dict[str, Any] = None) -> Dict[str, Any]:
    category = normalize_category(category) or "Bisnis"
    match_score = int(max(52, min(98, match_score)))
    paths = TOP_PATHS.get(category, [category])
    top_paths = [{"name": path, "score": max(50, match_score - i * 4)} for i, path in enumerate(paths[:3])]
    data = {
        "success": True,
        "prediction": category,
        "recommendedCategory": category,
        "recommended_career": category,
        "profileTitle": PROFILE_TITLES.get(category, category),
        "confidence": round(match_score / 100, 2),
        "matchScore": match_score,
        "level": "ELITE PROFILE" if match_score >= 85 else "STRONG PROFILE" if match_score >= 70 else "GROWING PROFILE",
        "topPathMatches": top_paths,
        "recommendedCareers": [p["name"] for p in top_paths],
        "matchedKeywords": matched_keywords,
        "reason": f"CAREVO membaca input dan menemukan kecocokan paling kuat ke {category}.",
        "suggestions": [
            f"Tambahkan skill dan project yang relevan dengan {category}.",
            "Lengkapi experience, education, skill, certification, dan project agar rekomendasi lebih akurat.",
            "Generate CV memakai data terbaru dari dashboard.",
        ],
    }
    if extra:
        data.update(extra)
    return data

@tf.keras.utils.register_keras_serializable()
class AttentionLayer(tf.keras.layers.Layer):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
    def build(self, input_shape):
        self.W = self.add_weight(name="attention_weight", shape=(input_shape[-1], 1), initializer="random_normal", trainable=True)
        super().build(input_shape)
    def call(self, inputs):
        score = tf.matmul(inputs, self.W)
        weights = tf.nn.softmax(score, axis=1)
        context = weights * inputs
        return tf.reduce_sum(context, axis=1)
    def get_config(self):
        return super().get_config()

@tf.keras.utils.register_keras_serializable()
class CareerModel(tf.keras.Model):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        class_count = max(1, len(labels) or 8)
        self.vectorizer = tf.keras.layers.TextVectorization(max_tokens=MAX_TOKENS, output_mode="int", output_sequence_length=MAX_LEN)
        self.embedding = tf.keras.layers.Embedding(MAX_TOKENS, EMBEDDING_DIM, mask_zero=True)
        self.bigru = tf.keras.layers.Bidirectional(tf.keras.layers.GRU(64, return_sequences=True, dropout=0.3))
        self.attention = AttentionLayer()
        self.dense = tf.keras.layers.Dense(128, activation="relu")
        self.dropout = tf.keras.layers.Dropout(0.5)
        self.output_layer = tf.keras.layers.Dense(class_count, activation="softmax")
    def call(self, inputs, training=False):
        x = self.vectorizer(inputs)
        x = self.embedding(x)
        x = self.bigru(x, training=training)
        x = self.attention(x)
        x = self.dense(x)
        x = self.dropout(x, training=training)
        return self.output_layer(x)
    def get_config(self):
        return super().get_config()

def model_predict(text: str) -> Dict[str, Any]:
    cleaned = clean_text(text)
    if model is None or label_encoder is None:
        return {"ranked": [], "engine": "rule-fallback-model-not-loaded"}
    try:
        pred = model(tf.constant([cleaned]), training=False).numpy()[0]
        class_labels = [str(item) for item in label_encoder.classes_]
        ranked = []
        for index, raw_label in enumerate(class_labels):
            label = normalize_category(raw_label)
            score = float(pred[index]) if index < len(pred) else 0.0
            rs = rule_scores(cleaned)
            ranked.append({"category": label, "rawScore": score, "matchedKeywords": rs["matched"].get(label, [])})
        ranked.sort(key=lambda item: item["rawScore"], reverse=True)
        return {"ranked": ranked, "engine": "tensorflow-career_model.keras"}
    except Exception as exc:
        return {"ranked": [], "engine": f"tensorflow-error-rule-fallback: {exc}"}

def normalize_result(text: str) -> Dict[str, Any]:
    cleaned = clean_text(text or "general profile")
    model_result = model_predict(cleaned)
    return hybrid_recommendation(cleaned, model_result=model_result, model_loaded=model is not None)

def add_value(parts: List[str], value: Any) -> None:
    if value is None:
        return
    if isinstance(value, (str, int, float)):
        v = str(value).strip()
        if v:
            parts.append(v)
        return
    if isinstance(value, list):
        for item in value:
            add_value(parts, item)
        return
    if isinstance(value, dict):
        for key in ["fullName", "careerInterest", "professionalTitle", "shortBio", "location", "name", "category", "level", "proficiency", "language", "jobTitle", "companyName", "employmentType", "description", "skillsUsed", "institutionName", "degreeMajor", "gpa", "title", "role", "status", "certificateName", "issuer"]:
            add_value(parts, value.get(key))

def collect_text_from_cv(cv_data: Dict[str, Any]) -> str:
    parts: List[str] = []
    for key in ["personalInfo", "profile", "skills", "skillList", "experiences", "educations", "projects", "projectList", "certifications", "certificationList", "languages", "languageList"]:
        add_value(parts, cv_data.get(key))
    return clean_text(" ".join(parts))

@app.on_event("startup")
def load_assets():
    global model, label_encoder, labels, load_error
    load_error = ""
    try:
        label_path = next((p for p in LABEL_PATHS if p.exists()), None)
        if label_path:
            with open(label_path, "rb") as f:
                label_encoder = pickle.load(f)
            labels = [str(item) for item in label_encoder.classes_]
        elif FALLBACK_LABELS_PATH.exists():
            labels = json.loads(FALLBACK_LABELS_PATH.read_text(encoding="utf-8"))
        else:
            labels = list(PROFILE_TITLES.keys())

        if MODEL_PATH.exists():
            custom_objects = {"AttentionLayer": AttentionLayer, "CareerModel": CareerModel}
            model = tf.keras.models.load_model(MODEL_PATH, custom_objects=custom_objects, compile=False)
        else:
            load_error = f"Model file tidak ditemukan: {MODEL_PATH}"
    except Exception as exc:
        model = None
        load_error = str(exc)

@app.get("/health")
def health():
    return {"success": True, "message": "CAREVO AI Service aktif.", "modelLoaded": model is not None, "labelLoaded": label_encoder is not None, "labels": labels, "modelPath": str(MODEL_PATH), "loadError": load_error}

@app.post("/predict")
def predict(payload: CareerRequest):
    return normalize_result(payload.skills_text)

@app.post("/predict-career")
def predict_career(payload: PredictCareerRequest):
    text = clean_text((payload.text or "") + " " + collect_text_from_cv(payload.cvData or {}))
    return normalize_result(text)

def add_section_title(doc, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(13)
    return p

def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(str(text or ""))
    run.font.size = Pt(10)
    return p

@app.post("/generate-ats-cv")
def generate_ats_cv(payload: GenerateCvRequest):
    if Document is None:
        return {"success": False, "message": "python-docx belum terinstall. Jalankan: pip install python-docx"}
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(str(payload.name or "CAREVO USER").upper())
    run.bold = True
    run.font.size = Pt(20)
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(" | ".join([x for x in [payload.phone, payload.email, payload.linkedin, payload.portfolio] if x]))
    loc = doc.add_paragraph()
    loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    loc.add_run(payload.location or "")
    add_section_title(doc, "Professional Summary")
    doc.add_paragraph(payload.professional_summary or "")
    add_section_title(doc, "Education")
    for edu in payload.education:
        p = doc.add_paragraph()
        p.add_run(str(edu.get("university") or edu.get("institutionName") or "Institution")).bold = True
        if edu.get("location"):
            p.add_run(f" | {edu.get('location')}")
        degree = doc.add_paragraph()
        degree.add_run(f"{edu.get('degree') or edu.get('degreeMajor') or ''} | GPA: {edu.get('gpa', '')} | {edu.get('year', '')}").italic = True
        for item in edu.get("details", []):
            add_bullet(doc, item)
    add_section_title(doc, "Work Experience")
    for exp in payload.experiences:
        p = doc.add_paragraph()
        p.add_run(str(exp.get("company") or exp.get("companyName") or "Company")).bold = True
        if exp.get("location"):
            p.add_run(f" | {exp.get('location')}")
        role = doc.add_paragraph()
        role.add_run(f"{exp.get('position') or exp.get('jobTitle') or ''} | {exp.get('period', '')}").italic = True
        for item in exp.get("responsibilities", []) or [exp.get("description", "")]:
            if item:
                add_bullet(doc, item)
    add_section_title(doc, "Projects")
    for project in payload.projects:
        p = doc.add_paragraph()
        p.add_run(str(project.get("title") or "Project")).bold = True
        doc.add_paragraph(str(project.get("description") or ""))
    add_section_title(doc, "Certifications")
    for cert in payload.certifications:
        add_bullet(doc, cert)
    add_section_title(doc, "Skills")
    doc.add_paragraph(payload.skills or "")
    add_section_title(doc, "Languages")
    doc.add_paragraph(payload.languages or "")
    out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    out.close()
    doc.save(out.name)
    return FileResponse(out.name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="CAREVO_AI_ATS_CV.docx")
